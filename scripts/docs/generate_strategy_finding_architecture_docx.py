"""Generate the strategy-finding pipeline architecture Word document."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Polygon
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "docs" / "exports"
DOCX_PATH = OUTPUT_DIR / "strategy_finding_pipeline_architecture_2026-07-27.docx"
FLOWCHART_PATH = OUTPUT_DIR / "strategy_finding_pipeline_flowchart_2026-07-27.png"
FONT_PATH = Path(r"C:\Windows\Fonts\msjh.ttc")

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "DCEAF7"
GREEN = "548235"
LIGHT_GREEN = "E2F0D9"
AMBER = "BF7A00"
LIGHT_AMBER = "FFF2CC"
PURPLE = "7030A0"
LIGHT_PURPLE = "E4DFEC"
RED = "C00000"
LIGHT_RED = "FCE4D6"
GRAY = "5B6573"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text: str,
    *,
    color: str = "222222",
    bold: bool = False,
    size: float = 9.5,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def shade_header(row, fill: str = NAVY) -> None:
    for cell in row.cells:
        set_cell_shading(cell, fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9, bold=True)
                run.font.color.rgb = rgb(WHITE)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    tail = paragraph.add_run(" 頁")
    set_run_font(tail, size=8)


def configure_section(section, *, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def configure_document(doc: Document) -> None:
    configure_section(doc.sections[0])
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Title", 28, NAVY),
        ("Subtitle", 14, GRAY),
        ("Heading 1", 19, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 11.5, GREEN),
    ):
        style = styles[name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(10)

    doc.core_properties.title = "策略 Finding Pipeline 現況架構與 ADR-0016 目標設計"
    doc.core_properties.subject = "GenAI discovery and deterministic strategy evaluation architecture"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "strategy finding, GenAI, backtest, ADR-0016, pipeline"
    doc.core_properties.comments = "Research-only architecture snapshot as of 2026-07-27."


def add_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = True
    paragraph = header.paragraphs[0]
    paragraph.text = "Strategy Finding Pipeline Architecture  |  2026-07-27"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        set_run_font(run, size=8)
        run.font.color.rgb = rgb(GRAY)
    footer = section.footer
    footer.is_linked_to_previous = True
    add_page_field(footer.paragraphs[0])


def add_callout(doc: Document, title: str, body: str, *, fill: str, accent: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, size=10.5, bold=True)
    run.font.color.rgb = rgb(accent)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(body)
    set_run_font(run, size=9.5)
    run.font.color.rgb = rgb("333333")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: Iterable[str], *, level: int = 0) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_run_font(run, size=10)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        set_run_font(run, size=10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, color=WHITE, bold=True, size=9)
    shade_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            fill = WHITE if row_index % 2 == 0 else LIGHT_GRAY
            set_cell_shading(cells[column_index], fill)
            set_cell_text(cells[column_index], value, size=8.7)
            if widths:
                cells[column_index].width = Inches(widths[column_index])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def flow_box(
    ax,
    x: float,
    y: float,
    width: float,
    height: float,
    text: str,
    *,
    face: str,
    edge: str,
    font: FontProperties,
    linestyle: str = "-",
    linewidth: float = 1.8,
    fontsize: float = 10,
):
    patch = FancyBboxPatch(
        (x, y),
        width,
        height,
        boxstyle="round,pad=0.035,rounding_size=0.12",
        facecolor=face,
        edgecolor=edge,
        linewidth=linewidth,
        linestyle=linestyle,
    )
    ax.add_patch(patch)
    ax.text(
        x + width / 2,
        y + height / 2,
        text,
        ha="center",
        va="center",
        fontproperties=font,
        fontsize=fontsize,
        color="#1F2937",
        linespacing=1.25,
    )
    return patch


def flow_arrow(
    ax,
    start: tuple[float, float],
    end: tuple[float, float],
    *,
    color: str = "#64748B",
    connectionstyle: str = "arc3",
    linestyle: str = "-",
    label: str | None = None,
    font: FontProperties | None = None,
    label_offset: tuple[float, float] = (0, 0),
):
    arrow = FancyArrowPatch(
        start,
        end,
        arrowstyle="-|>",
        mutation_scale=14,
        linewidth=1.5,
        color=color,
        linestyle=linestyle,
        connectionstyle=connectionstyle,
        shrinkA=2,
        shrinkB=2,
    )
    ax.add_patch(arrow)
    if label:
        ax.text(
            (start[0] + end[0]) / 2 + label_offset[0],
            (start[1] + end[1]) / 2 + label_offset[1],
            label,
            ha="center",
            va="center",
            fontproperties=font,
            fontsize=9,
            color=color,
            bbox={"facecolor": "white", "edgecolor": "none", "pad": 1.5},
        )


def create_flowchart(path: Path) -> None:
    font = FontProperties(fname=str(FONT_PATH)) if FONT_PATH.exists() else FontProperties()
    fig, ax = plt.subplots(figsize=(18.5, 10), dpi=190)
    ax.set_xlim(0, 20)
    ax.set_ylim(0, 11)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    ax.text(
        10,
        10.6,
        "策略 Finding Pipeline：目前架構與 ADR-0016 目標流程",
        ha="center",
        va="center",
        fontproperties=font,
        fontsize=20,
        fontweight="bold",
        color="#17365D",
    )
    ax.text(
        10,
        10.15,
        "一個使用者 prompt 可觸發多次 GenAI 呼叫；canonical evidence 只由 deterministic code 產生",
        ha="center",
        va="center",
        fontproperties=font,
        fontsize=10.5,
        color="#5B6573",
    )

    flow_box(
        ax,
        8.0,
        9.15,
        4.0,
        0.65,
        "使用者 prompt\n啟動一輪 strategy finding",
        face="#F2F4F7",
        edge="#5B6573",
        font=font,
        fontsize=10.5,
    )

    flow_box(
        ax,
        0.7,
        7.65,
        5.4,
        1.1,
        "新研究軌｜已存在、仍部分人工\narXiv / Semantic Scholar / Crossref\n→ fetch / firewall / scoring",
        face="#DCEAF7",
        edge="#2F75B5",
        font=font,
        fontsize=10.2,
    )
    flow_box(
        ax,
        0.7,
        6.15,
        5.4,
        1.0,
        "GenAI 解讀論文與可用資料\n→ 新機制 candidate spec\n現況：多數停在 pending_llm",
        face="#FFF2CC",
        edge="#BF7A00",
        font=font,
        fontsize=10,
    )
    flow_box(
        ax,
        13.9,
        7.65,
        5.4,
        1.1,
        "舊策略迭代軌｜已存在、仍部分人工\nSTRATEGY_HISTORY + Hypothesis / Experiment ledgers\n→ eligibility / K / retry 檢查",
        face="#DCEAF7",
        edge="#2F75B5",
        font=font,
        fontsize=9.8,
    )
    flow_box(
        ax,
        13.9,
        6.15,
        5.4,
        1.0,
        "GenAI / reviewer 提出 ex-ante 實質迭代\n不得依同輪結果追 gate\n現況：需 agent / human",
        face="#FFF2CC",
        edge="#BF7A00",
        font=font,
        fontsize=10,
    )

    flow_arrow(ax, (8.7, 9.15), (3.4, 8.75), connectionstyle="arc3,rad=0.18")
    flow_arrow(ax, (11.3, 9.15), (16.6, 8.75), connectionstyle="arc3,rad=-0.18")
    flow_arrow(ax, (3.4, 7.65), (3.4, 7.15))
    flow_arrow(ax, (16.6, 7.65), (16.6, 7.15))

    flow_box(
        ax,
        7.2,
        6.15,
        5.6,
        1.15,
        "候選池與 preflight\npaper provenance / family dedupe / data / timing /\ncost / trial-K / signal_ref-runner readiness",
        face="#F2F4F7",
        edge="#5B6573",
        font=font,
        fontsize=10,
    )
    flow_arrow(ax, (6.1, 6.65), (7.2, 6.65))
    flow_arrow(ax, (13.9, 6.65), (12.8, 6.65))

    diamond = Polygon(
        [(10, 5.75), (12.65, 4.75), (10, 3.75), (7.35, 4.75)],
        closed=True,
        facecolor="#FFF2CC",
        edgecolor="#BF7A00",
        linewidth=1.9,
    )
    ax.add_patch(diamond)
    ax.text(
        10,
        4.78,
        "10–15 個 execution-ready？\n新論文機制 >= 8\n舊策略迭代 >= 2",
        ha="center",
        va="center",
        fontproperties=font,
        fontsize=10.2,
        color="#1F2937",
        linespacing=1.25,
    )
    flow_arrow(ax, (10, 6.15), (10, 5.75))

    flow_box(
        ax,
        0.8,
        4.05,
        5.3,
        1.15,
        "不合格候選留在 full funnel\n重複 / 無資料 / 論文無法驗證 / contract 無效 /\n缺 runner → 補論文、補 candidate 或補實作",
        face="#FCE4D6",
        edge="#C00000",
        font=font,
        fontsize=9.6,
    )
    flow_arrow(
        ax,
        (7.35, 4.75),
        (6.1, 4.75),
        color="#C00000",
        label="否",
        font=font,
        label_offset=(0, 0.22),
    )
    flow_arrow(
        ax,
        (3.45, 5.2),
        (3.45, 6.15),
        color="#C00000",
        connectionstyle="arc3,rad=-0.12",
        linestyle="--",
        label="backfill",
        font=font,
        label_offset=(-0.7, 0),
    )

    flow_box(
        ax,
        7.25,
        2.65,
        5.5,
        0.9,
        "ADR-0016 目標｜result-blind round manifest\ncandidate / family / source / prompt / model / spec hashes",
        face="#E4DFEC",
        edge="#7030A0",
        font=font,
        linestyle="--",
        fontsize=10,
    )
    flow_arrow(
        ax,
        (10, 3.75),
        (10, 3.55),
        color="#548235",
        label="是",
        font=font,
        label_offset=(0.35, 0.05),
    )

    flow_box(
        ax,
        7.15,
        1.25,
        5.7,
        0.95,
        "Deterministic Stage 2｜registry 目前有限\nscreening backtest / returns + data / distinctness / cost / power",
        face="#E2F0D9",
        edge="#548235",
        font=font,
        fontsize=9.8,
    )
    flow_arrow(ax, (10, 2.65), (10, 2.2), color="#548235")

    decision = Polygon(
        [(14.9, 2.35), (16.15, 1.72), (14.9, 1.09), (13.65, 1.72)],
        closed=True,
        facecolor="#E2F0D9",
        edgecolor="#548235",
        linewidth=1.8,
    )
    ax.add_patch(decision)
    ax.text(
        14.9,
        1.72,
        "Stage 2\nPASS？",
        ha="center",
        va="center",
        fontproperties=font,
        fontsize=10,
        color="#1F2937",
    )
    flow_arrow(ax, (12.85, 1.72), (13.65, 1.72), color="#548235")

    flow_box(
        ax,
        16.75,
        2.45,
        2.65,
        0.9,
        "Stage 3（PASS only）\nWF / CPCV / DSR / PSR\nrobustness / costs",
        face="#E2F0D9",
        edge="#548235",
        font=font,
        fontsize=9.2,
    )
    flow_arrow(
        ax,
        (16.15, 1.95),
        (16.75, 2.7),
        color="#548235",
        label="PASS",
        font=font,
        label_offset=(0.05, 0.18),
    )

    flow_box(
        ax,
        16.2,
        0.45,
        3.5,
        0.95,
        "Deterministic canonical report\n完整 funnel + 每個 candidate PASS / FAIL\n僅研究結論，不代表 live readiness",
        face="#DCEAF7",
        edge="#2F75B5",
        font=font,
        fontsize=9.2,
    )
    flow_arrow(
        ax,
        (16.15, 1.48),
        (16.6, 1.2),
        color="#C00000",
        label="FAIL",
        font=font,
        label_offset=(0.2, -0.12),
    )
    flow_arrow(ax, (18.08, 2.45), (18.05, 1.4), color="#548235")

    legend_items = [
        ("#DCEAF7", "#2F75B5", "目前已存在"),
        ("#FFF2CC", "#BF7A00", "部分人工 / partial"),
        ("#E4DFEC", "#7030A0", "ADR-0016 目標、尚未完整實作"),
        ("#E2F0D9", "#548235", "deterministic evidence"),
        ("#FCE4D6", "#C00000", "拒絕 / 補位 / fail-closed"),
    ]
    x = 0.7
    for face, edge, label in legend_items:
        patch = FancyBboxPatch(
            (x, 0.35),
            0.35,
            0.22,
            boxstyle="round,pad=0.01,rounding_size=0.03",
            facecolor=face,
            edgecolor=edge,
            linewidth=1.2,
        )
        ax.add_patch(patch)
        ax.text(
            x + 0.45,
            0.46,
            label,
            ha="left",
            va="center",
            fontproperties=font,
            fontsize=8.2,
            color="#475569",
        )
        x += 3.55

    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("策略 Finding Pipeline")
    set_run_font(run, size=30, bold=True)
    run.font.color.rgb = rgb(NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("目前架構、GenAI 邊界與 ADR-0016 目標設計")
    set_run_font(run, size=18, bold=True)
    run.font.color.rgb = rgb(BLUE)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(10)
    run = date.add_run("架構快照日期：2026-07-27")
    set_run_font(run, size=11)
    run.font.color.rgb = rgb(GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for index, (label, value, fill) in enumerate(
        (
            ("規則", "完整輪次 10–15 個策略", LIGHT_BLUE),
            ("配比", "新論文 ≥ 8；舊迭代 ≥ 2", LIGHT_AMBER),
            ("證據權責", "Deterministic code", LIGHT_GREEN),
        )
    ):
        cell = table.cell(0, index)
        set_cell_shading(cell, fill)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=9, bold=True)
        r.font.color.rgb = rgb(GRAY)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=11, bold=True)
        r.font.color.rgb = rgb(NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=140, bottom=140)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_callout(
        doc,
        "重要狀態聲明",
        "本文件同時呈現目前已存在的能力與 ADR-0016 目標。"
        "目前尚未完成一鍵式 8/2/10 manifest 驗證、通用 candidate runner、"
        "manifest-hash resume 與 round-bound canonical report；因此現有指令只能產生 advisory／limited probe，"
        "不得宣稱已完成完整自動化輪次。",
        fill=LIGHT_RED,
        accent=RED,
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(26)
    run = paragraph.add_run("Research-only architecture document")
    set_run_font(run, size=10, bold=True)
    run.font.color.rgb = rgb(GRAY)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("不代表策略、demo、shadow 或 live readiness")
    set_run_font(run, size=9.5)
    run.font.color.rgb = rgb(RED)

    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)

    doc.add_heading("1. 執行摘要", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "目前 pipeline 已具備論文抓取與 prompt firewall、taxonomy／歷史候選產生、"
        "Stage-2／Stage-3 registry、順序式 resumable orchestrator，以及 deterministic funnel report。"
        "真正的瓶頸不是「能否找到很多論文」，而是「能否把候選轉成有資料、有 timing contract、"
        "有 registered runner 且能產生 terminal backtest evidence 的 strategy」。"
    )
    set_run_font(run, size=10.3)

    add_bullets(
        doc,
        [
            "完整一輪：在任何同輪結果可見前，凍結 10–15 個 unique、execution-ready 策略。",
            "組成：至少 8 個 verified-paper-backed 新機制，至少 2 個既有合格策略的 ex-ante 實質迭代。",
            "不計數：參數格點、rename、重複 family、無法驗證的論文、缺資料、invalid contract、缺 runner。",
            "執行：deterministic 程式跑所有 sealed candidates 的 Stage 2；Stage 3 只跑 Stage-2 PASS。",
            "報告：papers → ideas → rejects/backfills → Stage 2 → Stage 3 全 funnel 必須可 reconciliation。",
        ],
    )
    add_callout(
        doc,
        "一句話架構",
        "GenAI 負責擴張 discovery 與產生受約束的 candidate spec；一般程式負責驗證、回測、metrics、"
        "trial/K、PASS/FAIL 與 canonical report；人類負責研究優先級與是否進入下一個風險階段。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    doc.add_heading("2. 目前 vs. ADR-0016 目標", level=1)
    add_table(
        doc,
        ["層次", "目前狀態", "ADR-0016 目標"],
        [
            [
                "啟動方式",
                "多個 CLI／agent 手動銜接；沒有單一 prompt→完整 round。",
                "一個 user prompt 觸發 bounded GenAI calls，再交給 deterministic execution。",
            ],
            [
                "候選數量",
                "idea generator 只有 max=15，沒有 complete-round minimum／track quota。",
                "sealed executable slate 10–15；new_research ≥8、existing_iteration ≥2。",
            ],
            [
                "新論文軌",
                "arXiv／Semantic Scholar／Crossref fetch、firewall、scoring handoff 已有；多數 draft 為 pending_llm。",
                "跨來源 DOI/arXiv/title dedupe；GenAI 產生 schema-valid executable candidate spec。",
            ],
            [
                "舊策略迭代軌",
                "可讀 STRATEGY_HISTORY／ledgers，但選擇與 rationale 仍由 agent／reviewer 串接。",
                "自動列出 eligible families；ex-ante iteration rationale 與 K/retry provenance 一起凍結。",
            ],
            [
                "Execution contract",
                "Stage-2／Stage-3 registry 僅涵蓋少數既有 family；unknown/new 停在 awaiting implementation。",
                "每個 counted candidate 都要有 known signal_ref／runner、data/timing/cost/trial contract。",
            ],
            [
                "Resume",
                "有 candidate-level state，但沒有 manifest hash／input drift guard。",
                "state 綁定 sealed manifest hash；atomic state、candidate error isolation。",
            ],
            [
                "Report",
                "已有 family-wide funnel projection，但不能證明單一 frozen round 完整走完。",
                "per-round canonical report 對齊 manifest、state、每個 candidate terminal artifact。",
            ],
            [
                "UI／部署",
                "Research Ops 僅 read-only projection／有限本地工具；無 pipeline runner 或 promotion control。",
                "仍維持 research-only；本 ADR 不改 demo／shadow／live gates。",
            ],
        ],
        widths=[1.25, 2.75, 2.75],
    )

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(landscape, landscape=True)
    add_header_footer(landscape)
    heading = doc.add_heading("3. 端到端流程圖", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "顏色標示現況與目標；紫色虛線節點是 ADR-0016 已接受、但尚未完整實作的控制層。"
    )
    set_run_font(run, size=9.5)
    doc.add_picture(str(FLOWCHART_PATH), width=Inches(10.15))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("圖 1｜Strategy Finding Pipeline 現況與目標端到端流程")
    set_run_font(run, size=9, bold=True)
    run.font.color.rgb = rgb(GRAY)

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(portrait)
    add_header_footer(portrait)

    doc.add_heading("4. 分層架構與主要元件", level=1)
    component_rows = [
        [
            "1. Discovery",
            "research/crypto-alpha-lab paper ingestion；literature_keyword_scorer.py",
            "抓取論文、prompt firewall、metadata／abstract scoring handoff。",
            "部分完成",
        ],
        [
            "2. Idea synthesis",
            "pipeline_idea_generator.py；run_pipeline_literature_ideas.py",
            "taxonomy 與 literature drafts；讀 ledgers／feedback tags。",
            "部分完成",
        ],
        [
            "3. Research controls",
            "family minting、history／K、data／power screens",
            "dedupe、data feasibility、distinctness、cost、power、trial accounting。",
            "已存在但未整合 8/2/10",
        ],
        [
            "4. Orchestration",
            "pipeline_orchestrator.py；run_pipeline_orchestrator.py",
            "順序式執行、candidate state、Stage 2→Stage 3 stop rules。",
            "已存在、registry 有限",
        ],
        [
            "5. Stage 2",
            "pipeline_stage2_registry.py + family-specific probes",
            "資料／distinctness／cost／power 與部分 screening returns。",
            "已存在、非 generic",
        ],
        [
            "6. Stage 3",
            "pipeline_stage3_registry.py + strategy-specific runners",
            "fold-refit WF/CPCV、DSR/PSR、cost/robustness checkpoint。",
            "PASS-only、非 generic",
        ],
        [
            "7. Reporting",
            "run_pipeline_funnel_report.py；Research Funnel UI projection",
            "family evidence、ledger timeline、trial/K reconciliation。",
            "已存在；非 round-bound",
        ],
        [
            "8. Target control",
            "round manifest + signal_ref executable contract",
            "8/2/10 validation、hash freeze、resume drift guard、full reconciliation。",
            "尚未實作",
        ],
    ]
    add_table(
        doc,
        ["Layer", "主要檔案／元件", "責任", "狀態"],
        component_rows,
        widths=[1.0, 2.4, 2.6, 1.0],
    )

    doc.add_heading("5. 權責邊界", level=1)
    add_table(
        doc,
        ["角色", "可以做", "不能做／限制"],
        [
            [
                "GenAI",
                "擴展論文 query、解讀 verified paper、對照可用資料、提出新機制與舊策略迭代、輸出 schema-valid spec。",
                "manifest sealing 前看同輪 OOS/fold；直接執行 arbitrary generated code；計算或 override canonical gates；寫 canonical report。",
            ],
            [
                "Deterministic 程式",
                "驗證 schema／provenance／hash；跑 Stage 2／Stage 3；計算 metrics、trial/K、PASS/FAIL；輸出 artifacts/report。",
                "不得隱藏 rejected candidate、改寫 immutable evidence、用缺資料的替代來源、跳過 stop rules。",
            ],
            [
                "Codex implementation",
                "把 approved spec 接到 signal_ref／runner；補 timing/leakage test；執行 terminal commands；檢查 artifacts。",
                "不得依聊天記憶改策略假設；不得把 target doc 說成已實作；不得繞過 research/config authority。",
            ],
            [
                "Claude／reviewer",
                "論文與 novelty 審查、strategy/risk critique、lookahead／overfit／cost review。",
                "不得用 narrative 覆蓋 deterministic FAIL；不得把同輪結果導回 hidden retune。",
            ],
            [
                "使用者",
                "決定研究優先級、可接受的 limited probe、是否由 terminal／Codex 啟動 deterministic run，以及後續風險階段。",
                "完整輪次若少於十個仍只能標記 incomplete／limited；另需明確批准 demo／shadow／live。",
            ],
        ],
        widths=[1.1, 3.0, 3.0],
    )

    doc.add_heading("6. Candidate 與 Artifact Contract", level=1)
    doc.add_heading("6.1 Counted candidate 必備欄位", level=2)
    add_bullets(
        doc,
        [
            "Identity：candidate_id、hypothesis_id、family_id、track（new_research／existing_iteration）。",
            "Provenance：paper DOI／arXiv／URL 或舊策略 source hypothesis；retrieval、prompt、model、template、spec hash。",
            "Executable contract：signal_ref／runner、data source、timing/as-of rule、entry/exit、cost、grid、validation path。",
            "Research controls：family-cumulative n_trials、K used/limit、breadth、power inputs、distinctness reference。",
            "State：schema/data preflight PASS；不得是 pending_llm、unknown runner 或 invalid contract。",
        ],
    )

    doc.add_heading("6.2 主要 artifacts", level=2)
    add_table(
        doc,
        ["Artifact", "目前／目標", "用途"],
        [
            ["weekly paper screen／score handoff", "目前", "論文 metadata、摘要、filter／score 與 feedback。"],
            ["idea_batch.json／hypothesis draft", "目前", "候選與 skipped reasons；尚不代表 executable。"],
            ["stage2_feasibility.json", "目前", "data／distinctness／cost／power 等 terminal evidence。"],
            ["checkpoint1／Stage-3 artifacts", "目前", "WF/CPCV/DSR/PSR、cost/robustness 與 trial/K。"],
            ["round_manifest.json", "目標", "凍結 10–15 個 executable candidate、8/2 配比與所有 hashes。"],
            ["orchestrator_state.json + manifest hash", "目標強化", "resume、input drift guard、per-candidate terminal status。"],
            ["canonical round report", "目標", "papers→ideas→reject/backfill→Stage 2→Stage 3 的完整 reconciliation。"],
        ],
        widths=[2.1, 1.1, 3.8],
    )

    doc.add_heading("7. Gate 與決策邏輯", level=1)
    add_numbered(
        doc,
        [
            "Discovery oversample：先找比 final slate 更多的 verified papers／candidate specs；分開計算 paper、idea、executable strategy、parameter cells。",
            "Preflight：family/history dedupe、paper verification、data/timing/cost contract、runner readiness；不合格者留 funnel 並 backfill。",
            "Seal：只有 10–15 executable、new≥8、iteration≥2 才能 seal；seal 後才可讓同輪結果可見。",
            "Stage 2：每個 sealed candidate 都要產生 deterministic screening backtest／research-return terminal artifact；含 data、distinctness、cost、power。",
            "Stage 3：只跑 Stage-2 PASS；維持既有 fold-refit WF/CPCV、DSR/PSR、robustness、honest n_trials/K。",
            "Report：任一 candidate error 必須可見；不能因前幾個成功就把未完成 round 標記 completed。",
            "Research verdict：PASS 只表示通過該研究 gate；demo／shadow／live 仍須既有 gates 與使用者明確批准。",
        ],
    )

    add_callout(
        doc,
        "最低數量的真正含義",
        "十個不是十個參數 cell，也不是十張 idea 卡。十個代表十個 unique strategy contracts，"
        "各自具有 verified provenance、可用資料、registered runner，以及 deterministic terminal evaluation。",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )

    doc.add_page_break()
    doc.add_heading("8. 目前主要缺口與風險", level=1)
    add_table(
        doc,
        ["優先級", "缺口／風險", "影響", "建議控制"],
        [
            [
                "P0",
                "沒有 result-blind 8/2/10 manifest validator",
                "少量 candidates 仍可能被誤稱完整輪次。",
                "在 DB／backtest 前 fail-fast；limited_probe 必須 explicit。",
            ],
            [
                "P0",
                "Resume 未綁 manifest hash；state write／candidate error isolation 不完整",
                "輸入漂移、單一錯誤中斷後續、round 完整性無法證明。",
                "hash-bound resume、atomic state、per-candidate terminal error。",
            ],
            [
                "P1",
                "Literature 與 iteration entry points 分離；paper 跨來源無完整 dedupe",
                "無一鍵雙軌；同一 paper／mechanism 重複計數。",
                "DOI→arXiv→normalized title identity；unified candidate pool。",
            ],
            [
                "P1",
                "LLM scoring／spec handoff 仍需 session/manual",
                "規模與重現性受限。",
                "provider-neutral JSON adapter；凍結 model/prompt/template hashes。",
            ],
            [
                "P2",
                "Stage-2／Stage-3 registry 少且含歷史 identity assumptions",
                "大部分新 family 無法真正 backtest。",
                "沿用 signal_ref contract，逐個 mechanism 補 deterministic function＋timing/leakage check。",
            ],
            [
                "P3",
                "Current report 非 round-bound，可能混入歷史 family evidence",
                "無法證明本輪 10–15 全部完成。",
                "只從 sealed manifest＋本輪 state/artifacts 產生 canonical report。",
            ],
            [
                "Later",
                "Sequential execution／單一 DB connection",
                "10–15 策略可能較慢。",
                "先量測；只有 profiling 證明需要才加 bounded workers。",
            ],
            [
                "Always",
                "Quota filler、hallucinated citation、hidden retune",
                "錯誤但看似合理的 PASS。",
                "verified paper identity、novelty dedupe、result-blind seal、honest trial/K。",
            ],
        ],
        widths=[0.65, 2.0, 2.0, 2.4],
    )

    doc.add_heading("9. 最小自動化路線", level=1)
    roadmap = [
        ("Phase 0｜Round integrity", "round manifest schema；8/2/10 validator；unique IDs；executable readiness；manifest hash resume；per-round report skeleton。"),
        ("Phase 1｜Unified GenAI handoff", "整合論文與舊策略 iteration candidates；跨來源 paper dedupe；GenAI 只輸出 schema-valid JSON。"),
        ("Phase 2｜Executable breadth", "擴充 signal_ref registry／candidate-specific Stage-2 runner；每個新 mechanism 至少一個 timing/leakage self-check。"),
        ("Phase 3｜One prompt / one command", "prepare→seal→execute→report；atomic state；per-candidate exception isolation；resume。"),
        ("Phase 4｜Measured scaling", "先記錄各 candidate runtime／DB load，再決定是否加入 bounded parallel backtests。"),
    ]
    for title, body in roadmap:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(title)
        set_run_font(run, size=10.5, bold=True)
        run.font.color.rgb = rgb(BLUE)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        run = paragraph.add_run(body)
        set_run_font(run, size=9.8)

    add_callout(
        doc,
        "現階段不建議先做的事",
        "不要先導入 Airflow／多 agent framework／任意 Python codegen／大量 parallel workers。"
        "這些都不會解決目前最核心的 executable-contract 與 round-integrity 缺口；"
        "先讓順序式 10–15 strategy run 可被完整驗證與重現。",
        fill=LIGHT_GRAY,
        accent=GRAY,
    )

    doc.add_heading("10. 現在可採用的操作模式", level=1)
    add_bullets(
        doc,
        [
            "Agent-assisted discovery：使用者以一個 prompt 指定完整 round；Codex／Claude 可做多次 bounded literature search 與 spec synthesis。",
            "JSON handoff：candidate spec 經 schema／provenance／runner readiness review，再交 deterministic terminal command。",
            "Deterministic execution：可由使用者手動執行 terminal，或由 Codex 在授權範圍內執行。",
            "現況標籤：在 Phase 0–3 完成前，任何少於十個或含 missing runner 的 run 都只能叫 limited probe／incomplete round。",
        ],
    )

    doc.add_heading("11. Source of Truth 與檔案索引", level=1)
    add_table(
        doc,
        ["類型", "權威／主要檔案"],
        [
            ["架構決策", "docs/ADR/0016-genai-discovery-deterministic-strategy-evaluation.md"],
            ["完整輪次規則", "docs/AI_WORKFLOW.md；docs/DOMAIN_RULES.md R6.8／R6.9；docs/INVARIANTS.md I53／I54"],
            ["現況元件", "docs/FEATURE_MAP.md — Strategy Research Pipeline Automation"],
            ["資料／artifact 流", "docs/DATA_FLOW.md；scripts/run_pipeline_funnel_report.py"],
            ["策略歷史／trial-K", "docs/STRATEGY_HISTORY.md；docs/HYPOTHESIS_LEDGER.md；docs/EXPERIMENT_REGISTRY.md"],
            ["候選 executable contract", "docs/superpowers/specs/2026-06-30-drafted-candidate-stage3-contract.md"],
            ["主要程式", "backtesting/pipeline_idea_generator.py；pipeline_orchestrator.py；pipeline_stage2_registry.py；pipeline_stage3_registry.py"],
            ["現況缺口", "docs/KNOWN_ISSUES.md — ADR-0016/F56/F57/I53/I54"],
        ],
        widths=[1.45, 5.55],
    )

    doc.add_heading("12. 外部論文來源", level=1)
    sources = [
        ("arXiv API", "https://info.arxiv.org/help/api/index.html"),
        ("Semantic Scholar Academic Graph API", "https://api.semanticscholar.org/api-docs/graph"),
        ("Crossref REST API", "https://www.crossref.org/documentation/retrieve-metadata/rest-api/"),
    ]
    for name, url in sources:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(f"{name}: {url}")
        set_run_font(run, size=9.5)

    add_callout(
        doc,
        "文件結論",
        "目前架構已有 discovery、research controls、registered execution 與 reporting 的骨架，"
        "但尚未形成可一次完成 10–15 個 executable strategies 的閉環。最優先的不是再增加 agent，"
        "而是補上 result-blind manifest、通用 executable contract、hash-bound resume 與 round-bound report。",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_flowchart(FLOWCHART_PATH)
    document = build_document()
    document.save(DOCX_PATH)
    print(DOCX_PATH)
    print(FLOWCHART_PATH)


if __name__ == "__main__":
    main()
