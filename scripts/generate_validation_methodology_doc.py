"""Generate the strategy-validation methodology document (DOCX, Traditional Chinese).

Companion to ``scripts/generate_backtest_external_validation_report.py`` (the
slide deck). This script produces a complete, plain-language reference document
describing how the project currently validates backtest results by replaying
the same artifacts through third-party reference engines and comparing each
factor (signal / indicator / trade / pnl / source data).

Content source of truth: ``docs/validation_lab_report_zh.md`` and
``docs/ai_collaboration.md``. This is documentation only; it does not claim
live readiness and does not change any strategy, risk, config, or result
artifact.

Requires python-docx (``pip install python-docx``).
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path("docs/validation_methodology_zh.docx")
CJK_FONT = "Microsoft JhengHei"


def _apply_cjk(style) -> None:
    """Force a CJK-capable font (latin + eastAsia) on a paragraph/char style."""
    style.font.name = CJK_FONT
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), CJK_FONT)
    rfonts.set(qn("w:hAnsi"), CJK_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


def _set_fonts(doc: Document) -> None:
    for name in (
        "Normal",
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Bullet",
        "List Number",
    ):
        try:
            _apply_cjk(doc.styles[name])
        except KeyError:
            pass
    doc.styles["Normal"].font.size = Pt(11)


def h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def para(doc: Document, text: str, *, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def build(doc: Document) -> None:
    _set_fonts(doc)

    title = doc.add_heading("本專案策略驗證方法說明", level=0)
    title.alignment = 1
    sub = para(
        doc,
        "用第三方回測軟體交叉驗證回測結果——流程、因子拆解，與 vectorbt / backtrader / Nautilus 的分工",
    )
    sub.alignment = 1
    meta = para(doc, f"Quant Strategy 團隊｜版本日期 {date.today().isoformat()}｜內部文件")
    meta.alignment = 1

    # 1
    h(doc, "1. 文件目的與讀者")
    para(
        doc,
        "本文件說明本專案目前如何用第三方回測軟體驗證策略的回測結果，供內部團隊與覆核者閱讀。"
        "核心作法是：把我們自己跑出來的回測結果，丟給獨立的外部回測軟體，用同一批資料、同一組參數、"
        "同一套規則重算一遍，再逐項比對，確認結果一致且可以重現。",
    )
    para(doc, "範圍聲明：", bold=True)
    bullets(
        doc,
        [
            "本文件描述的是「回測可信度」的驗證，回答的是「同條件下計算結果是否一致、能否重現」。",
            "它不是策略會不會賺錢的證明，也不是上線交易的決策依據。",
            "策略是否上線（demo / shadow / live）另有獨立的審核流程，不在本文件範圍。",
        ],
    )

    # 2
    h(doc, "2. 名詞解釋")
    para(doc, "為避免專有名詞造成理解障礙，先用白話定義本文件會用到的詞：")
    table(
        doc,
        ["名詞", "白話解釋"],
        [
            ["回測（backtest）", "用歷史資料模擬策略在過去的表現。"],
            ["回測產物（artifact）", "一次回測存下來的檔案：價格、訊號、交易、權益曲線等。"],
            ["訊號（signal）", "策略決定進場 / 出場的時間點與方向。"],
            ["指標（indicator）", "由價格算出的數值，例如均線、EMA、MACD。"],
            ["重播 / replay", "本專案自己的回測引擎，會真的走風控、下單、成交。"],
            ["外部回測工具", "獨立的第三方 / 開源軟體，用來覆核我們算得對不對。"],
            ["嚴格關卡（gate）", "必須完全對齊才算通過的檢查。"],
            ["參考（advisory）", "會記錄差異供覆核，但不直接判定不通過。"],
            ["ct_val（合約乘數）", "把合約張數換算成實際部位大小的乘數，影響損益與風險計算。"],
            ["DB parity", "價格資料與資料庫標準資料是否逐筆同源的檢查。"],
        ],
    )

    # 3
    h(doc, "3. 為什麼要做外部驗證")
    para(doc, "問題", bold=True)
    para(
        doc,
        "我們自己的回測結果可信嗎？如果用自己寫的引擎、自己驗自己，缺乏客觀、獨立的第二意見，"
        "很難說服別人（或自己）結果沒有實作上的偏差。",
    )
    para(doc, "作法", bold=True)
    para(
        doc,
        "拿成熟的開源回測軟體，吃同一批資料、同一組參數、同一套規則，獨立重算一次，再逐項比對差異。"
        "等於請第二位裁判覆核同一場比賽。",
    )
    para(doc, "為什麼這樣有效", bold=True)
    bullets(
        doc,
        [
            "低成本：採用開源工具，本地就能跑，不依賴付費閉源平台。",
            "可重跑：條件固定，任何時候都能複現同樣的比對。",
            "可審計：每次比對都留下 JSON / CSV 證據，不靠口頭說明或聊天記錄。",
            "交叉比對：要兩套以上工具都對齊才算數，降低單一工具自身偏差的影響。",
        ],
    )

    # 4
    h(doc, "4. 整體驗證流程")
    para(doc, "驗證分成五個步驟，每一步都對應 repo 內可查的檔案：")
    numbered(
        doc,
        [
            "我們的回測：跑出回測產物（價格、訊號、交易、權益曲線），存放於 results/。",
            "拆成因子：把結果拆成可以分開驗的因子（訊號、指標、交易、績效、資料來源）。",
            "外部重算：用 vectorbt、backtrader、Nautilus 各自獨立重算對應的因子。",
            "逐因子比對：標記哪些對齊、哪些有差異、哪些差異需要修正。",
            "結論與證據：輸出通過 / 待補的結論，並保留 validation_result.json 與差異 CSV。",
        ],
    )
    para(
        doc,
        "因為每一步都落成檔案，整個驗證可以被重跑、被覆核，而不是依賴某次對話或某個人的記憶。",
    )

    # 5
    h(doc, "5. 把驗證拆成幾個因子（核心觀念）")
    para(
        doc,
        "驗證的關鍵觀念是：一次回測結果不是一個黑箱，而是可以拆成幾個獨立因子，每個因子分開驗、"
        "分開判定。這樣既能精準定位問題，也能清楚說明「目前到底驗到哪一層」。",
    )
    table(
        doc,
        ["因子", "驗什麼", "怎麼驗", "目前狀態"],
        [
            [
                "訊號（進出場點）",
                "同一根 K 棒、同方向、同動作是否一致",
                "外部工具用相同價格重算 crossover，比對時間 / 方向 / 動作",
                "已交叉驗證一致（嚴格關卡）",
            ],
            [
                "指標數據",
                "均線 / EMA / MACD 等數值是否算得相同",
                "外部工具重算指標值並與我們的值比對",
                "隨訊號一併比對，一致",
            ],
            [
                "交易（下單→成交）",
                "訊號變成下單與成交的過程",
                "外部工具產生對照用 trades",
                "參考性質（外部撮合 ≠ 本專案撮合）",
            ],
            [
                "績效（報酬 / 回撤）",
                "報酬率、最大回撤等績效數字",
                "外部 equity / metrics 對照",
                "參考性質",
            ],
            [
                "資料來源",
                "價格結構、合約乘數來源、與資料庫是否同源",
                "檢查 OHLCV 結構、ct_val 來源、選用的 DB parity",
                "結構與乘數已驗；資料庫對齊待補",
            ],
        ],
    )
    para(doc, "嚴格 vs 參考的差別：", bold=True)
    bullets(
        doc,
        [
            "嚴格關卡（訊號、指標）：必須完全對齊，有任何需要修正的差異就不算通過。",
            "參考（交易、績效）：差異會被記錄供覆核者引用，但不會自動判定整體不通過，"
            "因為外部工具的撮合與成本模型本來就和本專案不同。",
        ],
    )

    # 6
    h(doc, "6. 三套外部回測工具")
    h(doc, "6.1 vectorbt", level=2)
    para(doc, "定位：最快的「向量化」重算工具，一次把整段價格算完，適合大量、快速地比對訊號點。")
    bullets(
        doc,
        [
            "輸入：price_series.csv 的收盤序列、策略參數、用 signals.csv 比對。",
            "重算：warmup 後計算 SMA / EMA / MACD 與 crossover；用 Portfolio.from_signals 產生權益。",
            "設定：init_cash = 初始權益、fees = 0、freq 依 bar 對應年化頻率。",
            "能驗：指標數值、進出場訊號點（嚴格）。",
            "限制：不驗成交與撮合；費用設 0，產生的績效僅供參考。",
        ],
    )
    h(doc, "6.2 backtrader", level=2)
    para(doc, "定位：「逐根 K 棒」事件驅動工具，模擬時間一根一根往前走，比較貼近真實下單與決策流程。")
    bullets(
        doc,
        [
            "輸入：OHLCV 轉成 PandasData；bt.Cerebro 以初始權益 setcash。",
            "重算：TechnicalStrategy.next() 每根 K 棒更新指標狀態，發出市價單。",
            "參數：MA 20/50、EMA 20/50、MACD 12/26/9；warmup = slow（或 slow + signal）。",
            "能驗：訊號的時間點與方向（嚴格）。",
            "限制：使用市價單，撮合語意與本專案不同；下單 / 績效僅供參考。",
        ],
    )
    h(doc, "6.3 Nautilus", level=2)
    para(doc, "定位：目標是「高保真撮合」工具（含排隊、部分成交、資金費），但目前尚未啟用撮合引擎。")
    bullets(
        doc,
        [
            "目前 v1：把 artifact 訊號模擬成 long/flat 交易，輸出 reference_nautilus_*.csv 與 export_manifest.json。",
            "現況角色：engine_execution = not_run、reference_role = advisory；不能單獨讓可攜驗證關卡通過。",
            "目標路徑：若恢復 order-book 資料，建立 Nautilus catalog + L2/L3 order book + order/fill events + "
            "資金費 cashflows，對接撮合引擎，達成可獨立驗證的成交保真度。",
            "現況限制：standalone L2 runner 已停用，目前不維護 order-book 資料；排隊優先權與部分成交留待下一階段。",
        ],
    )
    h(doc, "6.4 本專案 replay（被驗證的對象）", level=2)
    para(
        doc,
        "本專案自己的回測重播（replay）引擎會真的走風控、下單、成交、出場，因此它是「被驗證的對象」，"
        "不是外部工具。上面三套外部工具的角色，是覆核 replay 算得對不對。",
    )

    # 7
    h(doc, "7. 工具能驗什麼、不能驗什麼（總表）")
    table(
        doc,
        ["工具", "能驗什麼（強項）", "限制 / 不能驗"],
        [
            ["vectorbt", "指標數值、進出場訊號點；速度快、可大量比對", "不驗成交與撮合；費用設 0，績效僅供參考"],
            ["backtrader", "訊號時間點與方向；逐根 K 棒模擬，貼近流程", "用市價單，撮合語意與本專案不同；績效僅供參考"],
            ["Nautilus", "（目標）撮合保真度、排隊、部分成交、資金費", "現況只做重播 / 匯出，尚未啟用撮合引擎，不能單獨當通過依據"],
            ["本專案 replay", "真實走風控、下單、成交、出場（被驗證對象）", "需要外部工具當第二意見才客觀"],
        ],
    )
    para(
        doc,
        "簡單說：vectorbt 與 backtrader 目前負責「訊號」這個最關鍵因子；Nautilus 負責未來的"
        "「成交保真度」；本專案 replay 是被覆核的對象。",
    )

    # 8
    h(doc, "8. 目前驗證成果")
    para(doc, "測試條件：", bold=True)
    bullets(
        doc,
        [
            "標的：Binance BTC 永續、1 小時 K 棒。",
            "區間：2024-01-01 ~ 2026-04-30，共 20,400 根 K 棒，資料覆蓋率 100%。",
            "策略與參數：MA（10/200）、EMA（10/200）、MACD（12/26/9）。",
        ],
    )
    para(doc, "結果：三個技術策略的進出場訊號，用 vectorbt 與 backtrader 兩套工具獨立重算後皆與我們的結果對齊。")
    table(
        doc,
        ["策略", "訊號數", "vectorbt", "backtrader", "比對結果"],
        [
            ["MA（10/200）", "228", "重算一致", "重算一致", "通過"],
            ["EMA（10/200）", "252", "重算一致", "重算一致", "通過"],
            ["MACD（12/26/9）", "1558", "重算一致", "重算一致", "通過"],
        ],
    )
    para(doc, "判讀重點：", bold=True)
    bullets(
        doc,
        [
            "代表在相同資料、參數、規則下，我們的訊號邏輯可以被外部工具完整重現，無需修正的差異為 0。",
            "本批使用 strategy_fill（理想成交）產物，且只選 vectorbt / backtrader；因此屬於訊號層的"
            "交叉驗證，admissibility = advisory_only。",
            "資料庫逐筆對齊（DB parity）目前為 SKIP；Nautilus 未納入本批。",
        ],
    )

    # 9
    h(doc, "9. 目前缺口與下一步")
    table(
        doc,
        ["缺口", "目前狀態", "下一步"],
        [
            ["資料庫對齊（Binance 1H）", "資料庫缺 Binance 1 小時標準 K 棒，DB parity 為 SKIP", "補資料（重新匯入或由 1 分鐘重採樣）後，開啟 DB parity 重跑"],
            ["撮合保真度（Nautilus）", "完整撮合引擎尚未啟用，僅做重播 / 匯出", "若恢復 order-book 資料，建立 catalog 與 L2/L3 撮合路徑"],
            ["實際成交率 / 成交模型", "realistic 成交下，訊號變成交的比例偏低", "決定小單 / 殘倉 / 排隊成交政策，或用 shadow / demo 校準"],
            ["績效穩健度", "尚未做樣本外 / walk-forward / CPCV", "策略要往前推進前必須補上"],
        ],
    )
    para(
        doc,
        "下一步順序：先補資料庫對齊，再提升撮合保真度，最後補績效穩健度檢查。這些都是把「回測可信」"
        "從訊號層往執行層與績效層延伸的後續工作，不影響目前訊號層已驗證的結論。",
    )

    # 10
    h(doc, "10. 驗證結果怎麼判讀")
    para(doc, "validation_result.json 的主要欄位與白話判讀：")
    table(
        doc,
        ["欄位", "白話判讀"],
        [
            ["status", "整體結果。通過只代表該批驗證範圍沒有硬性失敗，不等於可上線。"],
            ["admissibility", "若為 advisory_only，表示只能當參考證據，不能單獨作為推進依據。"],
            ["signal_logic_gate.passed", "最關鍵的嚴格關卡：至少一個 vectorbt / backtrader 訊號邏輯通過且需修正差異為 0。"],
            ["portable_validation_gate.passed", "是否具備合格的可攜外部驗證路徑；只做重播 / 匯出不算通過。"],
            ["source_data_validation.status", "資料與來源檢查：必要產物、OHLCV 結構、ct_val 來源是否符合。"],
            ["...checks.db_parity.status", "資料庫對齊是否執行。SKIP 不是失敗，但也不能宣稱已通過。"],
            ["...checks.ct_val_provenance.status", "合約乘數來源是否權威；會影響損益 / 風險判讀。"],
            ["actionable_mismatch_count", "需要修正或覆核的差異數。訊號邏輯中只要非 0 就會擋下。"],
        ],
    )

    # 11
    h(doc, "11. 附錄：操作入口與證據位置")
    para(doc, "CLI", bold=True)
    bullets(
        doc,
        [
            "python scripts/run_differential_validation.py --run-id <id> --strategy <name> "
            "--fixture-run-id <id> --engines vectorbt,backtrader,nautilus --validation-id <stable_id>",
        ],
    )
    para(doc, "API", bold=True)
    bullets(
        doc,
        [
            "POST /strategy-validation/run",
            "GET /strategy-validation/contracts",
            "POST /{run_id}/differential-validation/run",
            "GET artifact endpoints（取得 mismatch / reference 檔）",
        ],
    )
    para(doc, "UI", bold=True)
    bullets(
        doc,
        [
            "Validation view 顯示各引擎角色；訊號為嚴格，指標 / 交易 / 績效標示為參考。",
            "覆核者提示 PASS 的適用範圍；可預覽 mismatches_*.csv。",
        ],
    )
    para(doc, "證據檔案位置", bold=True)
    bullets(
        doc,
        [
            "results/validation_lab_*（各策略 run 的產物）",
            "results/.../validation/<validation_id>/validation_result.json 與 mismatches_*.csv",
            "results/engine_consistency_fixture/fixture_manifest.json（離線 smoke 用的凍結窗口）",
        ],
    )

    # 附錄 B：實測明細
    h(doc, "附錄 B：實測明細（2026-06-22 / 2026-06-23）")
    para(doc, "2026-06-22 訊號到下單檢查（初始風控 500 USD 單筆上限、部位上限 30%）：", bold=True)
    table(
        doc,
        ["策略", "參數", "訊號", "下單", "成交", "拒絕", "主要拒絕原因"],
        [
            ["MA", "10/200", "117", "5", "31", "112", "fat_finger"],
            ["EMA", "10/200", "127", "4", "22", "123", "fat_finger"],
            ["MACD", "12/26/9", "779", "779", "15", "0", "—"],
        ],
    )
    para(
        doc,
        "250/1.0 風控重跑（含 bounded reduce-only bypass）：MA 117/117/30/0、EMA 126/126/10/0、"
        "MACD 779/779/13/0。MA / EMA 的出場拒絕清零（原因是出場誤用了進場的 fat-finger 上限）；"
        "剩餘的低成交率來自 realistic 成交模型，不是訊號錯誤。",
    )
    para(doc, "Dual Output（MACD 全期，診斷用，非績效排名）：", bold=True)
    table(
        doc,
        ["設定", "訊號", "下單", "成交", "成交率", "報酬", "最大回撤"],
        [
            ["strategy_fill（理想成交）", "1558", "1558", "1558", "100%", "0.39%", "-6.99%"],
            ["realistic（真實成交）", "779", "779", "3 (+1 強平)", "0.39%", "0.66%", "-2.52%"],
        ],
    )
    para(
        doc,
        "Claude 2026-06-23 長區間（strategy_fill）：MA / EMA / MACD 在 vectorbt 與 backtrader 的訊號邏輯"
        "皆通過、需修正差異為 0、可攜驗證關卡為 true；DB parity = SKIP、admissibility = advisory_only。"
        "ct_val（BTC-USDT-SWAP）來源為資料庫、值 1.0、venue 為 binance、屬權威來源。",
    )

    para(doc, "")
    foot = para(
        doc,
        "來源：docs/validation_lab_report_zh.md、docs/ai_collaboration.md、"
        "backtesting/differential_validation.py。本文件為文件性質，不宣稱任何策略可上線交易。",
    )
    foot.runs[0].font.size = Pt(8)
    foot.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def write_doc(out: Path = OUT) -> Path:
    doc = Document()
    doc.core_properties.title = "本專案策略驗證方法說明"
    doc.core_properties.author = "Claude"
    build(doc)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


if __name__ == "__main__":
    path = write_doc()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
