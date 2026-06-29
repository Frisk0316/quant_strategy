"""Generate the Strategy Research Pipeline reference document (DOCX, Traditional Chinese).

Content source of truth:
``docs/superpowers/specs/2026-06-25-strategy-research-pipeline-design.md`` and the
companion manual chapter ``docs/manual/90-research-pipeline.md``. Documentation
only: it does not claim live readiness and does not change any strategy, risk,
config, gate, or result artifact.

Requires python-docx (``pip install python-docx``).
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path("docs/strategy_research_pipeline_zh.docx")
CJK_FONT = "Microsoft JhengHei"


def _apply_cjk(style) -> None:
    style.font.name = CJK_FONT
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), CJK_FONT)
    rfonts.set(qn("w:hAnsi"), CJK_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


def _set_fonts(doc: Document) -> None:
    for name in (
        "Normal", "Title", "Subtitle", "Heading 1", "Heading 2",
        "Heading 3", "List Bullet", "List Number",
    ):
        try:
            _apply_cjk(doc.styles[name])
        except KeyError:
            pass
    doc.styles["Normal"].font.size = Pt(11)


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, *, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def build(doc: Document) -> None:
    _set_fonts(doc)

    title = doc.add_heading("策略研究管線", level=0)
    title.alignment = 1
    sub = para(doc, "Strategy Research Pipeline — 按一次就跑到「通過 Gate 的短名單」的有界自動化研究流程")
    sub.alignment = 1
    meta = para(doc, f"Quant Strategy 團隊｜版本日期 {date.today().isoformat()}｜內部文件")
    meta.alignment = 1

    # 1
    h(doc, "1. 這是什麼")
    para(
        doc,
        "策略研究管線把「找文獻 → 評估可行性 → 做進回測系統 → 正式回測篩 Gate → 發布」"
        "接成有界自動化：使用者啟動一個預先登記的批次，按一次就跑到「通過 Gate 的短名單」，"
        "中間不用逐步下 prompt。重點是省人工，不是放寬標準。",
    )
    para(doc, "前置共識：", bold=True)
    bullets(
        doc,
        [
            "這是全專案 overfitting 風險最高的活動；Gate（DSR ≥ 0.95、PSR ≥ 0.95、honest n_trials、"
            "leak-free、differential validation、ct_val provenance）不可為了讓策略過而放鬆。",
            "「試幾次才會過」沒有固定答案：誠實計數下沒有真 edge 的想法期望通過數 ≈ 0；真正槓桿是先驗品質，不是試更多次。",
            "角色：Claude 做研究 / spec / 審查；Codex 實作 trading-core；發布是使用者的決定。",
        ],
    )
    para(doc, "設計來源：docs/superpowers/specs/2026-06-25-strategy-research-pipeline-design.md（status: draft）。")

    # 2
    h(doc, "2. 流程")
    para(doc, "使用者啟動批次 {candidates, runtime_cap, data_tier}，整批先寫進 ledger（過不過都算數），"
              "再由 driver（一個 Claude session）對每個候選依序執行：")
    numbered(
        doc,
        [
            "Stage 1 文獻→假設（研究 subagent）→ 寫 HYPOTHESIS_LEDGER（H-xxx）。",
            "Stage 2 可行性檢查（研究 subagent）→ 不過就 skip 並記原因。",
            "Stage 3 實作 + 回測（Codex subagent）→ leak 回歸測試 + 兩段式回測 + Gate 證據 artifact。",
            "檢查點① Claude 證據審查 → 通過進短名單；沒過記 refuted / shelved。",
            "檢查點② 使用者決定是否發布。",
        ],
    )
    para(doc, "兩個人為關卡：① Claude 審回測證據，② 使用者決定發布；Stage 1→3 之間無人值守。")

    # 3
    h(doc, "3. 各階段契約")
    table(
        doc,
        ["階段", "負責", "產出 / 通過條件"],
        [
            ["Stage 1 文獻→假設", "研究 subagent",
             "HYPOTHESIS_LEDGER 假設（含 family、可測 signal/entry/exit/sizing/execution/risk spec、grid、validation path、預登記 n_trials 預算）"],
            ["Stage 2 可行性", "研究 subagent",
             "資料可得性、相關性上限、成本後 edge 嗅探；PASS 進 Stage 3，FAIL 則 skip 並記原因（0 trials）"],
            ["Stage 3 實作+回測", "Codex subagent",
             "trading-core 實作 + 強制 leak 回歸測試 + 兩段式回測 + 機器可讀 Gate 證據 artifact"],
            ["檢查點①", "Claude",
             "用 REVIEW_QUESTIONS.md / CRITIQUE_PROTOCOL.md + 部署 Gate 條文審證據"],
            ["檢查點② / 發布", "使用者",
             "通過者接進系統成 enabled:false 已驗證候選 + ledger 標 supported"],
        ],
    )

    # 4
    h(doc, "4. 關鍵規則")
    bullets(
        doc,
        [
            "編排：單 session + subagent-driven-development skill；不是新框架、不是 cron。",
            "n_trials 按假設家族(family)累計：family 的 n_trials = 該 family 歷來所有批次的 grid 組合數 + 重試次數總和，"
            "餵進 CPCV 算 DSR。防「跨 batch 拆小」「同 family 微調重試」兩種灌水——把計數變嚴格，不是放寬。",
            "重試上限 K（預設 2）：同一經濟機制重試 K 次仍不過 → shelve 並升級給使用者；不會無限調參。",
            "重試 vs 新 family：一次 attempt 內掃 grid 是正常搜索（全數計入 family）；同機制調旋鈕重跑算重試（吃 K）；"
            "真的不同經濟機制才是新 family（K 歸零）。把重試 relabel 成新 family 由檢查點① 守住。",
            "兩段式回測：parquet research-tier 預篩 → DB venue-scoped CPCV 正式。",
            "發布定義：只接成 enabled:false 已驗證候選 + ledger 標 supported；不自動上線、不碰任何 demo/shadow/live Gate。",
            "記錄載體：長期真相寫 HYPOTHESIS_LEDGER.md + EXPERIMENT_REGISTRY.md；results/<batch_id>/ 的 JSON 與短名單是可丟 scratch，檢查點時對帳同步回 ledger。",
        ],
    )

    # 5
    h(doc, "5. 檢查點① Claude 證據審查清單")
    para(doc, "直接重用 REVIEW_QUESTIONS.md / CRITIQUE_PROTOCOL.md + ai_collaboration.md 的 Gate 條文，不另造 checklist：")
    numbered(
        doc,
        [
            "n_trials 誠實：傳進 CPCV 的 == ledger family 累計（不是程式寫死）。",
            "leak-free：leak 回歸測試存在且通過 + spot-check lag 邏輯。",
            "DSR 不變量：DSR ≤ PSR(0) 成立、由修正後 harness 算。",
            "idealized-fill 排除：idealized_fill == false。",
            "differential validation：portable_validation_gate 通過或誠實標 blocked。",
            "ct_val provenance：全 authoritative 且 venue 一致。",
            "門檻：DSR ≥ 0.95 且 PSR ≥ 0.95。",
            "裁決：supported → 進短名單；否則 refuted / shelved，n_trials 照記進 family。",
            "重試 / 新 family 判定：確認這個 attempt 是重試（吃 K）還是新 family（K 歸零），防 relabel 繞過上限。",
        ],
    )

    # 6
    h(doc, "6. 目前狀態（截至本文件版本日期）")
    bullets(
        doc,
        [
            "Stage 1 machinery（driver / templates、family 累計 n_trials、invariant I23、change manifest）已建好；"
            "完整的自動 driver 編排尚未跑過。",
            "第一批候選順序 [S7, S5, S6] 的 Stage 3 回測已單獨跑出 refit artifacts，全部未通過 Gate。",
            "目前沒有任何候選通過 Gate，無 promotion / live / demo / shadow。",
        ],
    )
    para(doc, "第一批 refit 結果：", bold=True)
    table(
        doc,
        ["候選", "WF OOS Sharpe", "CPCV OOS Sharpe", "DSR", "PSR", "裁決"],
        [
            ["S6 time-series momentum", "0.0088", "0.5422", "0.1963", "0.7387", "statistical_gate_passed:false"],
            ["S7 basis mean-reversion", "-0.4359", "-1.1124", "≈0", "≈0", "shelved_pending_research_review"],
            ["S5 residual mean-reversion", "—", "—", "—", "—", "nonzero_grid_activity:false（資料-universe artifact，非支持也非反駁）"],
        ],
    )

    # 7
    h(doc, "7. 分階段路線圖")
    table(
        doc,
        ["階段", "內容", "解鎖條件"],
        [
            ["Stage 1（現在）", "單 session、現有 backlog、一個檢查點、手動啟動", "—"],
            ["Stage 2", "開背景平行 agent", "把檢查點① 的 leak / n_trials / DSR 不變量檢查沉澱成自動 invariant/測試"],
            ["Stage 3", "排程 cron/loop 自動執行 + 文獻搜索 ingestion 加品質過濾", "Stage 2 穩定"],
        ],
    )

    para(doc, "")
    foot = para(
        doc,
        "來源：docs/superpowers/specs/2026-06-25-strategy-research-pipeline-design.md、"
        "docs/manual/90-research-pipeline.md、docs/ai_collaboration.md。"
        "本文件為文件性質，不宣稱任何策略可上線交易；發布權在使用者。",
    )
    foot.runs[0].font.size = Pt(8)
    foot.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def write_doc(out: Path = OUT) -> Path:
    doc = Document()
    doc.core_properties.title = "策略研究管線"
    doc.core_properties.author = "Claude"
    build(doc)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


if __name__ == "__main__":
    path = write_doc()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
